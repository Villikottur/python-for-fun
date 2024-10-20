from urllib.request import urlopen, Request
from bs4 import BeautifulSoup
from docx import Document

#scrapes the current rate
url = "https://www.exchange-rates.org/converter/usd-eur"
headers = {'User-Agent':'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_2) AppleWebKit/601.3.9 (KHTML, like Gecko) Version/9.0.2 Safari/601.3.9'}
req = Request(url=url, headers=headers)
html = urlopen(req).read()
soup = BeautifulSoup(html, "html.parser")
results = soup.find("span", class_="rate-to").text
rate = float(results[:-4])

#opens the doc
print("Tell me the name of the invoice file (must be in the same folder of this script):")
filename = input()
file = Document(filename+".docx")

#parses through the tables
for index, item in enumerate(file.tables):
    rows_n = len(file.tables[index].rows)
    col_n = len(file.tables[index].columns)   
    for x in range(rows_n):
        for y in range(col_n):
            if 'USD' in file.tables[index].cell(x,y).text:
                for i in range(len(file.tables[index].cell(x,y).paragraphs)):
                    #gets values and converts them
                    if file.tables[index].cell(x,y).paragraphs[i].text != "":
                        celp_txt = file.tables[index].cell(x,y).paragraphs[i].text
                        usd = celp_txt[celp_txt.find("")+len(""):celp_txt.rfind(" USD")]
                        usdv = float(usd.replace(",","."))
                        eur = round(usdv*rate,2)
                        #saves celp properties
                        fname = file.tables[index].cell(x,y).paragraphs[i].runs[0].font.name
                        fsize = file.tables[index].cell(x,y).paragraphs[i].runs[0].font.size
                        fbold = file.tables[index].cell(x,y).paragraphs[i].runs[0].font.bold
                        fital = file.tables[index].cell(x,y).paragraphs[i].runs[0].font.italic
                        funder = file.tables[index].cell(x,y).paragraphs[i].runs[0].font.underline
                        fstrike = file.tables[index].cell(x,y).paragraphs[i].runs[0].font.strike
                        fcol = file.tables[index].cell(x,y).paragraphs[i].runs[0].font.color.rgb
                        fhigh = file.tables[index].cell(x,y).paragraphs[i].runs[0].font.highlight_color
                        #clears celp and replaces it
                        file.tables[index].cell(x,y).paragraphs[i].clear()
                        file.tables[index].cell(x,y).paragraphs[i].add_run(str(usd)+" USD ("+str(eur).replace(".",",")+" EUR)")
                        file.tables[index].cell(x,y).paragraphs[i].runs[0].font.name = fname
                        file.tables[index].cell(x,y).paragraphs[i].runs[0].font.size = fsize
                        file.tables[index].cell(x,y).paragraphs[i].runs[0].font.bold = fbold 
                        file.tables[index].cell(x,y).paragraphs[i].runs[0].font.italic = fital
                        file.tables[index].cell(x,y).paragraphs[i].runs[0].font.underline = funder
                        file.tables[index].cell(x,y).paragraphs[i].runs[0].font.strike = fstrike
                        file.tables[index].cell(x,y).paragraphs[i].runs[0].font.color.rgb = fcol
                        file.tables[index].cell(x,y).paragraphs[i].runs[0].font.highlight_color = fhigh

#saves file and says bye                   
file.save(filename+".docx")
print("Hope you like this little program, thank you for testing it!")